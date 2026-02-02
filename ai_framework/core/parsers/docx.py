"""
DOCX document parser.

Handles Microsoft Word .docx files using python-docx.

Story 7.6: Document Parsing & Chunking
"""

import logging
from typing import Optional

from .base import DocumentParser, ParsedDocument, ParserError

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """
    Parser for Microsoft Word documents (.docx).
    
    Uses python-docx to extract text and metadata.
    """
    
    SUPPORTED_MIME_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",  # Legacy .doc (not supported, just listed)
    }
    
    SUPPORTED_EXTENSIONS = {
        ".docx",
    }
    
    def __init__(self):
        """Initialize DOCX parser."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are installed."""
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
    
    def can_parse(self, mime_type: str, file_extension: str) -> bool:
        """Check if this parser can handle the file type."""
        return (
            mime_type in self.SUPPORTED_MIME_TYPES or
            file_extension.lower() in self.SUPPORTED_EXTENSIONS
        )
    
    def parse(self, file_content: bytes, filename: str = "") -> ParsedDocument:
        """
        Parse DOCX content.
        
        Args:
            file_content: Raw DOCX bytes
            filename: Original filename
            
        Returns:
            ParsedDocument: Parsed DOCX document
            
        Raises:
            ParserError: If DOCX parsing fails
        """
        import io
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = self.docx.Document(docx_file)
            
            # Extract metadata
            metadata = {
                "filename": filename,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }
            
            # Add core properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            
            # Count sections/pages (approximation)
            section_count = len(doc.sections)
            
            return ParsedDocument(
                text=text,
                metadata=metadata,
                page_count=section_count,  # Approximation
            )
        
        except Exception as e:
            raise ParserError(f"Failed to parse DOCX {filename}: {str(e)}")
