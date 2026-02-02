"""
Document parser implementations for various formats.

Supports: PDF, DOCX, TXT, HTML, Markdown
"""

import io
import logging
from typing import BinaryIO

from .base import DocumentParser, ParsingError

logger = logging.getLogger(__name__)


class PDFParser(DocumentParser):
    """Parser for PDF documents using PyPDF2."""
    
    def __init__(self):
        try:
            import PyPDF2
            self.PyPDF2 = PyPDF2
        except ImportError:
            raise ParsingError(
                "PyPDF2 not installed. Install with: pip install PyPDF2"
            )
    
    def parse(self, file_stream: BinaryIO) -> str:
        """Extract text from PDF."""
        try:
            reader = self.PyPDF2.PdfReader(file_stream)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
            
            text = '\n\n'.join(text_parts)
            return self.clean_text(text)
            
        except Exception as e:
            raise ParsingError(f"PDF parsing failed: {e}")
    
    def supports_mime_type(self, mime_type: str) -> bool:
        """Check if PDF MIME type."""
        return mime_type in ['application/pdf', 'application/x-pdf']


class DOCXParser(DocumentParser):
    """Parser for Microsoft Word documents using python-docx."""
    
    def __init__(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ParsingError(
                "python-docx not installed. Install with: pip install python-docx"
            )
    
    def parse(self, file_stream: BinaryIO) -> str:
        """Extract text from DOCX."""
        try:
            doc = self.docx.Document(file_stream)
            
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = '\n\n'.join(text_parts)
            return self.clean_text(text)
            
        except Exception as e:
            raise ParsingError(f"DOCX parsing failed: {e}")
    
    def supports_mime_type(self, mime_type: str) -> bool:
        """Check if DOCX MIME type."""
        return mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]


class TextParser(DocumentParser):
    """Parser for plain text files."""
    
    def parse(self, file_stream: BinaryIO) -> str:
        """Extract text from TXT file."""
        try:
            content = file_stream.read()
            
            # Try UTF-8 first, fallback to latin-1
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1', errors='ignore')
            
            return self.clean_text(text)
            
        except Exception as e:
            raise ParsingError(f"Text parsing failed: {e}")
    
    def supports_mime_type(self, mime_type: str) -> bool:
        """Check if text MIME type."""
        return mime_type in [
            'text/plain',
            'text/markdown',
            'text/csv',
            'application/json',
            'application/xml',
            'text/html'
        ]


class GoogleDocsParser(DocumentParser):
    """
    Parser for Google Docs (exported as text).
    
    Note: Google Docs need to be exported to a supported format
    via the Drive API before parsing.
    """
    
    def parse(self, file_stream: BinaryIO) -> str:
        """Parse Google Docs (exported as plain text)."""
        # Google Docs are exported as plain text
        text_parser = TextParser()
        return text_parser.parse(file_stream)
    
    def supports_mime_type(self, mime_type: str) -> bool:
        """Check if Google Docs MIME type."""
        return mime_type == 'application/vnd.google-apps.document'
